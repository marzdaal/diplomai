from __future__ import annotations

import asyncio
import json
import uuid
from io import BytesIO
from typing import AsyncGenerator, List

from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel, Field

app = FastAPI(title="DiplomAI API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"] ,
    allow_headers=["*"]
)


def _make_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:8]}"


class GenerateRequest(BaseModel):
    topic: str = Field(min_length=3, max_length=240)


class ItemRequest(BaseModel):
    type: str


class TitlePageData(BaseModel):
    university: str = Field(min_length=2, max_length=200)
    faculty: str = Field(min_length=2, max_length=200)
    department: str = Field(min_length=2, max_length=200)
    work_title: str = Field(min_length=3, max_length=240, alias="workTitle")
    student_name: str = Field(min_length=2, max_length=120, alias="studentName")
    supervisor_name: str = Field(min_length=2, max_length=120, alias="supervisorName")
    city: str = Field(min_length=2, max_length=120)
    year: str = Field(min_length=4, max_length=10)

    model_config = {"populate_by_name": True}


class DocumentRequest(BaseModel):
    profile: str = Field(pattern="^(gost|msu|hse)$")
    title: TitlePageData
    goals: List[str] = Field(min_length=1)
    tasks: List[str] = Field(min_length=1)


async def _stream_items(topic: str) -> AsyncGenerator[bytes, None]:
    goals = [
        f"Определить теоретические основания исследования темы «{topic}».",
        f"Сформулировать методологию оценки и анализа по теме «{topic}».",
        f"Обосновать практическую значимость исследования по теме «{topic}»."
    ]
    tasks = [
        f"Собрать и систематизировать источники по теме «{topic}».",
        "Выявить ключевые подходы и определить критерии анализа.",
        "Сформировать набор показателей для оценки результатов.",
        "Проанализировать полученные данные и интерпретировать выводы.",
        "Сформулировать рекомендации по итогам исследования."
    ]

    for text in goals:
        payload = {"type": "goal", "id": _make_id("g"), "text": text}
        yield (json.dumps(payload, ensure_ascii=False) + "\n").encode("utf-8")
        await asyncio.sleep(0.4)

    for text in tasks:
        payload = {"type": "task", "id": _make_id("t"), "text": text}
        yield (json.dumps(payload, ensure_ascii=False) + "\n").encode("utf-8")
        await asyncio.sleep(0.35)


@app.post("/generate")
async def generate(request: GenerateRequest) -> StreamingResponse:
    return StreamingResponse(_stream_items(request.topic), media_type="application/json")


@app.post("/items/{item_id}/rephrase")
async def rephrase(item_id: str, request: ItemRequest) -> dict:
    return {
        "id": item_id,
        "text": f"Переформулированный вариант ({request.type}) для элемента {item_id}."
    }


@app.post("/items/{item_id}/replace")
async def replace(item_id: str, request: ItemRequest) -> dict:
    return {
        "id": item_id,
        "text": f"Новая альтернатива ({request.type}) для элемента {item_id}."
    }


def _build_outline(profile: str) -> list[str]:
    if profile == "msu":
        return [
            "Введение",
            "Глава 1. Теоретико-методологические основания",
            "Глава 2. Аналитическая часть",
            "Глава 3. Проектные решения и рекомендации",
            "Заключение",
            "Список литературы",
            "Приложения"
        ]
    if profile == "hse":
        return [
            "Введение",
            "Раздел 1. Теоретический обзор",
            "Раздел 2. Методология и анализ данных",
            "Раздел 3. Практическая часть и результаты",
            "Заключение",
            "Список литературы",
            "Приложения"
        ]
    return [
        "Введение",
        "Глава 1. Теоретические основы",
        "Глава 2. Аналитическая часть",
        "Глава 3. Практическая часть",
        "Заключение",
        "Список литературы",
        "Приложения"
    ]


def _add_centered(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold


def _generate_docx(payload: DocumentRequest) -> bytes:
    document = Document()

    _add_centered(document, payload.title.university, bold=True)
    _add_centered(document, payload.title.faculty)
    _add_centered(document, payload.title.department)
    document.add_paragraph()
    _add_centered(document, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", bold=True)
    _add_centered(document, payload.title.work_title, bold=True)
    document.add_paragraph()

    left = document.add_paragraph(f"Студент: {payload.title.student_name}")
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = document.add_paragraph(f"Научный руководитель: {payload.title.supervisor_name}")
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph()

    _add_centered(document, f"{payload.title.city} — {payload.title.year} г.")

    document.add_page_break()
    document.add_heading("Содержание", level=1)

    outline = _build_outline(payload.profile)
    for item in outline:
        document.add_paragraph(item)

    document.add_paragraph()
    document.add_heading("Цель работы", level=2)
    for goal in payload.goals:
        document.add_paragraph(goal, style="List Bullet")

    document.add_heading("Задачи исследования", level=2)
    for task in payload.tasks:
        document.add_paragraph(task, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


@app.post("/documents/generate")
async def generate_document(request: DocumentRequest) -> Response:
    content = _generate_docx(request)
    headers = {"Content-Disposition": "attachment; filename=diplom-structure.docx"}
    return Response(
        content,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers=headers
    )
